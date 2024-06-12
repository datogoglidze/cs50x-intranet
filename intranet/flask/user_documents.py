import datetime
import os
from uuid import uuid4

import pythoncom
from dependency_injector.wiring import Provide, inject
from docx import Document
from docx2pdf import convert
from flask import (
    Blueprint,
    redirect,
    render_template,
    request,
    send_from_directory,
    session,
)
from werkzeug import Response

from intranet.core.user_details import UserDetailsRepository
from intranet.core.user_documents import UserDocument
from intranet.error import apology, login_required
from intranet.flask.dependable import Container

user_documents = Blueprint(
    "user_documents", __name__, template_folder="../front/templates"
)


@user_documents.get("/documents")
@inject
@login_required
def user_details_page(
    details: UserDetailsRepository = Provide[Container.user_details_repository],
) -> str:
    files_and_dirs = os.listdir("vacations")
    last_name = details.read(session["user_id"]).last_name
    documents = [file for file in files_and_dirs if last_name in file]

    return render_template("user_documents.html", documents=documents)


@user_documents.get("/documents/<filename>")
@login_required
def pdf_viewer(filename: str) -> Response:
    return send_from_directory("../../vacations", filename)


@user_documents.post("/documents")
@inject
@login_required
def create_document(
    details: UserDetailsRepository = Provide[Container.user_details_repository],
) -> Response | tuple[str, int]:
    user_details = details.read(session["user_id"])
    user_document = UserDocument(
        id=str(uuid4()),
        first_name=user_details.first_name,
        last_name=user_details.last_name,
        dates=request.form.get("dates", ""),
    )

    if not user_document.first_name or not user_document.last_name:
        return apology("must fill details", 403)

    if not user_document.dates:
        return apology("must specify date", 403)

    generate_document(
        user_document.id,
        user_document.first_name,
        user_document.last_name,
        user_document.dates,
    )

    return redirect("/user-details")


def generate_document(_id: str, first_name: str, last_name: str, dates: str) -> None:
    document = Document("document_templates/vacation_template.docx")
    document.styles["Normal"].font.name = "Sylfaen"

    for paragraph in document.paragraphs:
        paragraph.text = paragraph.text.replace(
            "!<<DOC_ID>>",
            _id,
        )
        paragraph.text = paragraph.text.replace(
            "!<<FIRST_NAME>>",
            first_name,
        )
        paragraph.text = paragraph.text.replace(
            "!<<LAST_NAME>>",
            last_name[:-1] + "ის",
        )
        paragraph.text = paragraph.text.replace(
            "!<<DATE>>",
            dates,
        )

    document_name = f"{datetime.datetime.now().date()}-{last_name}-{first_name}"
    document.save(f"vacations/{document_name}.docx")

    pythoncom.CoInitialize()
    convert(
        f"vacations/{document_name}.docx",
        "vacations/",
    )
    os.remove(f"vacations/{document_name}.docx")
