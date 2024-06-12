import datetime
import os
from uuid import uuid4

import docx2pdf
import pythoncom
from dependency_injector.wiring import Provide, inject
from docx import Document
from flask import (
    Blueprint,
    redirect,
    render_template,
    request,
    send_from_directory,
    session,
)
from werkzeug import Response

from intranet.core.user_details import UserDetailsRepository
from intranet.core.user_documents import UserDocument
from intranet.error import apology, login_required
from intranet.flask.dependable import Container

documents = Blueprint("documents", __name__, template_folder="../front/templates")


@documents.get("/documents")
@inject
@login_required
def user_details_page(
    details: UserDetailsRepository = Provide[Container.user_details_repository],
) -> str:
    _documents = os.listdir("documents")
    last_name = details.read(session["user_id"]).last_name
    user_documents = [file for file in _documents if last_name in file]

    return render_template("user_documents.html", documents=user_documents)


@documents.get("/documents/<filename>")
@login_required
def pdf_viewer(filename: str) -> Response:
    return send_from_directory("../../documents", filename)


@documents.post("/documents")
@inject
@login_required
def create_document(
    details: UserDetailsRepository = Provide[Container.user_details_repository],
) -> Response | tuple[str, int]:
    user = details.read(session["user_id"])
    document = UserDocument(
        id=str(uuid4()),
        first_name=user.first_name,
        last_name=user.last_name,
        dates=request.form.get("dates", ""),
    )

    if not document.first_name or not document.last_name:
        return apology("must fill details", 403)

    if not document.dates:
        return apology("must specify date", 403)

    generate_document_with(
        document.id,
        document.first_name,
        document.last_name,
        document.dates,
    )

    return redirect("/documents")


def generate_document_with(
    document_id: str,
    first_name: str,
    last_name: str,
    dates: str,
) -> None:
    document = Document("document_templates/vacation_template.docx")
    document.styles["Normal"].font.name = "Sylfaen"

    for paragraph in document.paragraphs:
        paragraph.text = paragraph.text.replace(
            "!<<DOC_ID>>",
            document_id,
        )
        paragraph.text = paragraph.text.replace(
            "!<<FIRST_NAME>>",
            first_name,
        )
        paragraph.text = paragraph.text.replace(
            "!<<LAST_NAME>>",
            last_name[:-1] + "ის",
        )
        paragraph.text = paragraph.text.replace(
            "!<<DATE>>",
            dates,
        )

    new_document = (
        f"{datetime.datetime.now().strftime('%Y-%m-%d-%H-%M-%S')}"
        f"-{last_name}"
        f"-{first_name}"
    )
    document.save(f"documents/{new_document}.docx")

    pythoncom.CoInitialize()
    docx2pdf.convert(
        f"documents/{new_document}.docx",
        "documents/",
    )
    os.remove(f"documents/{new_document}.docx")
